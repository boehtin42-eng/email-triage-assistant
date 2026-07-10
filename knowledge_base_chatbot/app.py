import io
import html
import json
import re
import time
from typing import Any, Dict, List, Mapping, Optional, Tuple

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from pypdf import PdfReader

try:
    from google.oauth2.service_account import Credentials
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseDownload
except ImportError:
    Credentials = None
    build = None
    MediaIoBaseDownload = None

try:
    import google.generativeai as genai
except ImportError:
    genai = None


CHUNK_SIZE = 1200
CHUNK_OVERLAP = 200
MAX_CONTEXT_CHUNKS = 5
GOOGLE_DRIVE_SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]
GOOGLE_NATIVE_MIME_TYPES = {
    "application/vnd.google-apps.document",
    "application/vnd.google-apps.spreadsheet",
    "application/vnd.google-apps.presentation",
}
ANSWER_LANGUAGES = {
    "Burmese": {
        "instruction": "Always answer in Burmese/Myanmar language.",
        "audience": "Burmese-speaking staff member",
    },
    "English": {
        "instruction": "Always answer in clear English.",
        "audience": "English-speaking staff member",
    },
    "German": {
        "instruction": "Always answer in clear German.",
        "audience": "German-speaking staff member",
    },
}


st.set_page_config(page_title="Knowledge Base Chatbot", layout="wide")

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []


def get_secret(name: str) -> Optional[str]:
    try:
        value = st.secrets.get(name)
    except Exception:
        return None
    return str(value) if value else None


def get_service_account_info() -> Dict[str, Any]:
    try:
        raw_value = st.secrets.get("GOOGLE_SERVICE_ACCOUNT_JSON")
    except Exception as exc:
        raise RuntimeError("Could not read GOOGLE_SERVICE_ACCOUNT_JSON from Streamlit Secrets.") from exc

    if not raw_value:
        raise RuntimeError("GOOGLE_SERVICE_ACCOUNT_JSON is missing from Streamlit Secrets.")

    if isinstance(raw_value, Mapping):
        return dict(raw_value)

    try:
        return json.loads(str(raw_value))
    except json.JSONDecodeError as exc:
        raise RuntimeError(
            "GOOGLE_SERVICE_ACCOUNT_JSON is not valid JSON. Use triple single quotes in "
            "Streamlit Secrets and paste the full service account JSON."
        ) from exc


def clean_text(text: str) -> str:
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf(file) -> str:
    reader = PdfReader(file)
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"[Page {index}]\n{page_text}")
    return "\n\n".join(pages)


def extract_docx(file) -> str:
    document = Document(file)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_rows = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))
    return "\n".join(paragraphs + table_rows)


def extract_spreadsheet(file, file_name: str) -> str:
    if file_name.lower().endswith(".csv"):
        dataframe = pd.read_csv(file, dtype=str).fillna("")
        return dataframe.to_csv(index=False)

    sheets = pd.read_excel(file, sheet_name=None, dtype=str)
    parts = []
    for sheet_name, dataframe in sheets.items():
        dataframe = dataframe.fillna("")
        parts.append(f"[Sheet: {sheet_name}]\n{dataframe.to_csv(index=False)}")
    return "\n\n".join(parts)


def extract_text_from_bytes(file_name: str, data: bytes) -> str:
    extension = file_name.lower().rsplit(".", 1)[-1]

    if extension == "pdf":
        return clean_text(extract_pdf(io.BytesIO(data)))
    if extension == "docx":
        return clean_text(extract_docx(io.BytesIO(data)))
    if extension in {"xlsx", "xls", "csv"}:
        return clean_text(extract_spreadsheet(io.BytesIO(data), file_name))
    if extension in {"txt", "md"}:
        return clean_text(data.decode("utf-8", errors="ignore"))

    raise ValueError(f"Unsupported file type: {extension}")


def extract_text(uploaded_file) -> str:
    file_name = uploaded_file.name
    data = uploaded_file.getvalue()
    return extract_text_from_bytes(file_name, data)


@st.cache_resource
def get_drive_service():
    if Credentials is None or build is None:
        raise RuntimeError("Google Drive dependencies are not installed.")

    credentials = Credentials.from_service_account_info(
        get_service_account_info(),
        scopes=GOOGLE_DRIVE_SCOPES,
    )
    return build("drive", "v3", credentials=credentials, cache_discovery=False)


def is_transient_google_error(error: Exception) -> bool:
    message = str(error).lower()
    transient_markers = [
        "ssl",
        "record layer failure",
        "connection reset",
        "temporarily unavailable",
        "timeout",
        "timed out",
        "rate limit",
        "backend error",
    ]
    return any(marker in message for marker in transient_markers)


def execute_google_request(request, label: str, retries: int = 3):
    last_error = None
    for attempt in range(1, retries + 1):
        try:
            return request.execute()
        except Exception as exc:
            last_error = exc
            if attempt == retries or not is_transient_google_error(exc):
                break
            time.sleep(attempt)

    raise RuntimeError(f"{label} failed after {retries} attempt(s): {last_error}") from last_error


def download_drive_file(service, file_id: str) -> bytes:
    request = service.files().get_media(fileId=file_id)
    buffer = io.BytesIO()
    downloader = MediaIoBaseDownload(buffer, request)
    done = False
    attempts = 0
    while not done:
        try:
            _, done = downloader.next_chunk()
        except Exception as exc:
            attempts += 1
            if attempts >= 3 or not is_transient_google_error(exc):
                raise RuntimeError(f"Google Drive file download failed after 3 attempt(s): {exc}") from exc
            time.sleep(attempts)
    return buffer.getvalue()


def export_drive_file(service, file_id: str, mime_type: str) -> Tuple[str, bytes]:
    if mime_type == "application/vnd.google-apps.document":
        request = service.files().export(fileId=file_id, mimeType="text/plain")
        return ".txt", execute_google_request(request, "Google Docs export")
    if mime_type == "application/vnd.google-apps.spreadsheet":
        request = service.files().export(
            fileId=file_id,
            mimeType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        data = execute_google_request(request, "Google Sheets export")
        return ".xlsx", data
    if mime_type == "application/vnd.google-apps.presentation":
        request = service.files().export(fileId=file_id, mimeType="text/plain")
        return ".txt", execute_google_request(request, "Google Slides export")
    raise ValueError(f"Unsupported Google native file type: {mime_type}")


def load_drive_documents(folder_id: str) -> Tuple[List[Dict[str, str]], List[Dict[str, Any]], List[str]]:
    service = get_drive_service()
    query = f"'{folder_id}' in parents and trashed = false"
    request = service.files().list(
        q=query,
        fields="files(id,name,mimeType,size,modifiedTime)",
        orderBy="modifiedTime desc",
        pageSize=50,
        supportsAllDrives=True,
        includeItemsFromAllDrives=True,
    )
    response = execute_google_request(request, "Google Drive folder list")

    chunks: List[Dict[str, str]] = []
    summaries: List[Dict[str, Any]] = []
    errors: List[str] = []

    for file_info in response.get("files", []):
        file_id = file_info["id"]
        file_name = file_info["name"]
        mime_type = file_info.get("mimeType", "")
        try:
            if mime_type in GOOGLE_NATIVE_MIME_TYPES:
                suffix, data = export_drive_file(service, file_id, mime_type)
                parse_name = f"{file_name}{suffix}"
            else:
                data = download_drive_file(service, file_id)
                parse_name = file_name

            text = extract_text_from_bytes(parse_name, data)
            file_chunks = chunk_text(file_name, text)
            chunks.extend(file_chunks)
            summaries.append(
                {
                    "Source": "Google Drive",
                    "File": file_name,
                    "Type": "Google native" if mime_type in GOOGLE_NATIVE_MIME_TYPES else parse_name.rsplit(".", 1)[-1].upper(),
                    "Size KB": round(len(data) / 1024, 1),
                    "Chunks": len(file_chunks),
                }
            )
        except Exception as exc:
            errors.append(f"{file_name}: {exc}")

    return chunks, summaries, errors


def google_drive_enabled() -> bool:
    return bool(get_secret("GOOGLE_DRIVE_FOLDER_ID") and get_secret("GOOGLE_SERVICE_ACCOUNT_JSON"))


def build_upload_summary(uploaded_file, chunks: List[Dict[str, str]]) -> Dict[str, Any]:
    file_type = uploaded_file.name.rsplit(".", 1)[-1].upper()
    return {
        "Source": "Upload",
        "File": uploaded_file.name,
        "Type": file_type,
        "Size KB": round(len(uploaded_file.getvalue()) / 1024, 1),
        "Chunks": len(chunks),
    }


def chunk_text(source: str, text: str) -> List[Dict[str, str]]:
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + CHUNK_SIZE, len(text))
        chunk = text[start:end].strip()
        if chunk:
            chunks.append({"source": source, "text": chunk})
        if end == len(text):
            break
        start = max(0, end - CHUNK_OVERLAP)
    return chunks


def render_copy_button(text: str, button_label: str = "Copy answer") -> None:
    escaped_text = html.escape(text)
    escaped_label = html.escape(button_label)
    components.html(
        f"""
        <button
            onclick="navigator.clipboard.writeText(document.getElementById('copy-text').innerText);
                     this.innerText='Copied';"
            style="
                background:#ff4b4b;
                color:white;
                border:none;
                border-radius:6px;
                padding:0.55rem 0.8rem;
                font-weight:600;
                cursor:pointer;
            "
        >
            {escaped_label}
        </button>
        <pre id="copy-text" style="display:none;">{escaped_text}</pre>
        """,
        height=48,
    )


def tokenize(text: str) -> List[str]:
    return re.findall(r"[\w]+", text.lower())


def score_chunk(question_tokens: List[str], chunk: Dict[str, str]) -> int:
    chunk_text_lower = chunk["text"].lower()
    score = 0
    for token in question_tokens:
        if len(token) < 3:
            continue
        score += chunk_text_lower.count(token)
    return score


def find_relevant_chunks(question: str, chunks: List[Dict[str, str]]) -> List[Dict[str, str]]:
    question_tokens = tokenize(question)
    scored: List[Tuple[int, Dict[str, str]]] = []
    for chunk in chunks:
        score = score_chunk(question_tokens, chunk)
        if score > 0:
            scored.append((score, chunk))

    scored.sort(key=lambda item: item[0], reverse=True)
    return [chunk for _, chunk in scored[:MAX_CONTEXT_CHUNKS]]


def build_prompt(question: str, relevant_chunks: List[Dict[str, str]], answer_language: str) -> str:
    context = "\n\n".join(
        f"Source: {chunk['source']}\nContent:\n{chunk['text']}" for chunk in relevant_chunks
    )
    language_config = ANSWER_LANGUAGES[answer_language]
    return f"""
You are a company knowledge base assistant.
Answer only from the provided company documents.
If the documents do not contain the answer, say that the answer is not available in the uploaded documents.
{language_config["instruction"]}
Keep the answer clear, short, and practical for a {language_config["audience"]}.
Mention the source file names used.
Keep product names, company names, prices, dates, and source file names exactly as written in the documents.

Question:
{question}

Company document context:
{context}
""".strip()


def answer_with_gemini(
    question: str,
    relevant_chunks: List[Dict[str, str]],
    answer_language: str,
) -> Tuple[str, str]:
    api_key = get_secret("GEMINI_API_KEY")
    if not api_key or genai is None:
        return "", "No Gemini API key found."

    model_name = get_secret("GEMINI_MODEL") or "gemini-2.5-flash"
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        response = model.generate_content(build_prompt(question, relevant_chunks, answer_language))
        return (response.text or "").strip(), ""
    except Exception as exc:
        return "", f"Gemini answer generation failed: {exc}"


def render_source_chunks(chunks: List[Dict[str, str]]) -> None:
    st.subheader("Source snippets")
    for index, chunk in enumerate(chunks, start=1):
        with st.expander(f"{index}. {chunk['source']}"):
            st.write(chunk["text"])


def render_uploaded_docs_summary(document_summaries: List[Dict[str, Any]]) -> None:
    st.subheader("Uploaded documents")
    st.dataframe(pd.DataFrame(document_summaries), use_container_width=True, hide_index=True)


def render_chat_history() -> None:
    if not st.session_state.chat_history:
        return

    st.subheader("Chat history")
    if st.button("Clear chat history"):
        st.session_state.chat_history = []
        st.rerun()

    for index, item in enumerate(reversed(st.session_state.chat_history), start=1):
        with st.expander(f"{index}. {item['question']}"):
            st.caption(f"Language: {item['language']} | Sources: {', '.join(item['sources'])}")
            st.write(item["answer"])
            render_copy_button(item["answer"], "Copy this answer")


st.title("Internal Knowledge Base Chatbot")
st.caption("Upload company docs, ask questions, and get answers grounded in the uploaded files.")

with st.sidebar:
    st.header("Document Source")
    document_source = st.radio(
        "Use documents from",
        ["Upload only", "Google Drive folder", "Upload + Google Drive"],
        index=0,
    )

    st.header("Upload Docs")
    uploaded_files = []
    if document_source in {"Upload only", "Upload + Google Drive"}:
        uploaded_files = st.file_uploader(
            "PDF, DOCX, TXT, MD, CSV, XLSX",
            type=["pdf", "docx", "txt", "md", "csv", "xlsx", "xls"],
            accept_multiple_files=True,
        )

    if document_source in {"Google Drive folder", "Upload + Google Drive"}:
        st.header("Google Drive")
        if google_drive_enabled():
            st.success("Google Drive folder secrets found.")
            if st.button("Refresh Google Drive docs"):
                get_drive_service.clear()
                st.rerun()
        else:
            st.warning("Add GOOGLE_DRIVE_FOLDER_ID and GOOGLE_SERVICE_ACCOUNT_JSON in Streamlit Secrets.")

    answer_language = st.selectbox(
        "Answer language",
        list(ANSWER_LANGUAGES.keys()),
        index=0,
    )
    st.info("The app does not train a model. It searches the uploaded docs during this session.")

all_chunks: List[Dict[str, str]] = []
document_summaries: List[Dict[str, Any]] = []
extraction_errors = []

if uploaded_files:
    for uploaded_file in uploaded_files:
        try:
            text = extract_text(uploaded_file)
            file_chunks = chunk_text(uploaded_file.name, text)
            all_chunks.extend(file_chunks)
            document_summaries.append(build_upload_summary(uploaded_file, file_chunks))
        except Exception as exc:
            extraction_errors.append(f"{uploaded_file.name}: {exc}")

if document_source in {"Google Drive folder", "Upload + Google Drive"}:
    folder_id = get_secret("GOOGLE_DRIVE_FOLDER_ID")
    if google_drive_enabled() and folder_id:
        try:
            drive_chunks, drive_summaries, drive_errors = load_drive_documents(folder_id)
            all_chunks.extend(drive_chunks)
            document_summaries.extend(drive_summaries)
            extraction_errors.extend(drive_errors)
        except Exception as exc:
            extraction_errors.append(f"Google Drive: {exc}")

if extraction_errors:
    for error in extraction_errors:
        st.error(error)

if not all_chunks:
    st.warning("Add at least one company document from upload or Google Drive to start.")
    st.stop()

st.success(f"Loaded {len(document_summaries)} file(s) and created {len(all_chunks)} searchable text chunk(s).")
render_uploaded_docs_summary(document_summaries)

question = st.text_input(
    "Ask a question",
    placeholder="Example: What is our cancellation policy?",
)

if st.button("Ask", type="primary"):
    if not question.strip():
        st.error("Please enter a question.")
        st.stop()

    relevant_chunks = find_relevant_chunks(question, all_chunks)
    if not relevant_chunks:
        st.warning("No relevant text found in the uploaded documents.")
        st.stop()

    with st.spinner("Searching documents and preparing answer..."):
        answer, answer_error = answer_with_gemini(question, relevant_chunks, answer_language)

    st.subheader("Answer")
    if answer:
        st.write(answer)
        render_copy_button(answer)
        st.session_state.chat_history.append(
            {
                "question": question.strip(),
                "answer": answer,
                "language": answer_language,
                "sources": sorted({chunk["source"] for chunk in relevant_chunks}),
            }
        )
    else:
        st.warning(answer_error)
        st.info("The app is showing the most relevant source snippets instead.")

    render_source_chunks(relevant_chunks)

render_chat_history()
