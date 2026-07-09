from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors


OUTPUT_DIR = Path(__file__).parent
DOCX_PATH = OUTPUT_DIR / "sample_service_manual.docx"
PDF_PATH = OUTPUT_DIR / "sample_service_manual.pdf"


SECTIONS = [
    (
        "Company Overview",
        [
            "Company Name: Bright Cleaning Service",
            "Bright Cleaning Service provides home cleaning, office cleaning, deep cleaning, move-out cleaning, and monthly cleaning packages.",
            "The company focuses on reliable appointments, polite customer support, and clear follow-up after every service.",
        ],
    ),
    (
        "Business Hours",
        [
            "Our business hours are Monday to Saturday, 9:00 AM to 6:00 PM.",
            "We are closed on Sunday and public holidays.",
            "Urgent customer complaints should still be recorded and reviewed on the next working day.",
        ],
    ),
    (
        "Booking Rules",
        [
            "Normal cleaning should be booked at least 24 hours before the preferred appointment time.",
            "Deep cleaning should be booked at least 48 hours in advance.",
            "Staff should confirm the customer name, phone number, address, service type, preferred date, and preferred time before creating the booking.",
        ],
    ),
    (
        "Cancellation Policy",
        [
            "Customers can cancel or reschedule an appointment for free if they contact us at least 24 hours before the appointment.",
            "If a customer cancels less than 24 hours before the appointment, a cancellation fee may apply.",
            "If the company needs to reschedule because of staff availability, no cancellation fee should be charged.",
        ],
    ),
    (
        "Payment Policy",
        [
            "Customers can pay by cash, bank transfer, or online payment.",
            "Payment is usually collected after the service is completed.",
            "For monthly cleaning packages, payment can be made monthly in advance.",
        ],
    ),
    (
        "Complaint Handling",
        [
            "When a customer complains, staff should respond politely and ask for the service date, customer name, and problem details.",
            "Urgent complaints should be marked as high priority and sent to the manager.",
            "If the company made a clear service mistake, the customer may receive a partial refund or a free correction visit.",
        ],
    ),
    (
        "Staff Checklist",
        [
            "Arrive on time and wear a clean uniform.",
            "Speak politely and confirm the requested service before starting.",
            "Take before and after photos when required.",
            "Contact the office before accepting extra work that was not included in the original booking.",
        ],
    ),
]


FAQ_ROWS = [
    ("Do you provide cleaning supplies?", "Yes. We can bring cleaning supplies, or use the customer's preferred products."),
    ("Can I book for tomorrow?", "Yes, if there is availability. Normal cleaning should be booked 24 hours in advance."),
    ("How much is the monthly package?", "The price depends on the property size, number of rooms, and cleaning frequency."),
    ("What if I need to cancel?", "Cancel at least 24 hours before the appointment to avoid a cancellation fee."),
]


def add_docx_paragraph(document: Document, text: str, style: str = "Normal"):
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    return paragraph, run


def build_docx() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    for heading_name, size in [("Heading 1", 20), ("Heading 2", 16)]:
        styles[heading_name].font.name = "Arial"
        styles[heading_name].font.size = Pt(size)
        styles[heading_name].font.bold = False

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("Bright Cleaning Service - Sample Knowledge Base")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(26)

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run("Test document for Internal Knowledge Base Chatbot upload.")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(11)

    for heading, bullets in SECTIONS:
        document.add_heading(heading, level=1)
        for item in bullets:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(item)

    document.add_heading("Common Customer Questions", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Question"
    header_cells[1].text = "Approved Answer"
    for question, answer in FAQ_ROWS:
        cells = table.add_row().cells
        cells[0].text = question
        cells[1].text = answer

    document.add_heading("Suggested Chatbot Test Questions", level=1)
    for question in [
        "What is the cancellation policy?",
        "How should staff handle a complaint?",
        "Can a customer book deep cleaning tomorrow?",
        "How can customers pay?",
        "What should staff do before accepting extra work?",
    ]:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(question)

    document.save(DOCX_PATH)


def build_pdf() -> None:
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="BodyCustom",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="HeadingCustom",
            parent=styles["Heading1"],
            fontName="Helvetica",
            fontSize=15,
            leading=18,
            spaceBefore=12,
            spaceAfter=6,
        )
    )

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )

    story = [
        Paragraph("Bright Cleaning Service - Sample Knowledge Base", styles["Title"]),
        Paragraph("Test document for Internal Knowledge Base Chatbot upload.", styles["BodyCustom"]),
        Spacer(1, 0.15 * inch),
    ]

    for heading, bullets in SECTIONS:
        story.append(Paragraph(heading, styles["HeadingCustom"]))
        for item in bullets:
            story.append(Paragraph(f"- {item}", styles["BodyCustom"]))

    story.append(Paragraph("Common Customer Questions", styles["HeadingCustom"]))
    table_data = [
        [Paragraph("Question", styles["BodyCustom"]), Paragraph("Approved Answer", styles["BodyCustom"])]
    ]
    for question, answer in FAQ_ROWS:
        table_data.append(
            [
                Paragraph(question, styles["BodyCustom"]),
                Paragraph(answer, styles["BodyCustom"]),
            ]
        )
    table = Table(table_data, colWidths=[2.25 * inch, 4.25 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D0D5DD")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(table)

    story.append(Paragraph("Suggested Chatbot Test Questions", styles["HeadingCustom"]))
    for question in [
        "What is the cancellation policy?",
        "How should staff handle a complaint?",
        "Can a customer book deep cleaning tomorrow?",
        "How can customers pay?",
        "What should staff do before accepting extra work?",
    ]:
        story.append(Paragraph(f"- {question}", styles["BodyCustom"]))

    doc.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
